import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Title
title = doc.add_heading('BÁO CÁO: ỨNG DỤNG SPRING AI VÀO DỰ ÁN INTELIPATH', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Section 1
doc.add_heading('1. Giới thiệu: Template Spring AI là gì và ta đã làm gì?', level=1)
doc.add_paragraph('Trong quá trình phát triển dự án InteliPath, chúng ta đã quyết định tích hợp Spring AI (được truyền cảm hứng từ spring-ai-examples). Đây là một framework mới của hệ sinh thái Spring giúp kết nối các mô hình trí tuệ nhân tạo (LLM như OpenAI GPT-4o, Gemini) vào ứng dụng Java một cách cực kỳ chuẩn mực và dễ dàng.')
doc.add_paragraph('Thay vì phải gọi API REST trần trụi (dùng RestTemplate hoặc WebClient) tới OpenAI và tự xử lý các JSON phản hồi phức tạp, Spring AI cung cấp các Interface cấp cao (như ChatClient, VectorStore, Advisor) để xây dựng một AI Mentor toàn diện.')

# Section 2
doc.add_heading('2. Chúng ta đã làm được những gì với Spring AI?', level=1)
doc.add_paragraph('Nhờ áp dụng Spring AI, dự án InteliPath đã sở hữu một AI Virtual Mentor cực kỳ thông minh với các tính năng vượt trội:')
doc.add_paragraph('1. Sliding Window Memory (Bộ nhớ hội thoại): AI có khả năng nhớ lại 10 tin nhắn gần nhất của người dùng nhờ kỹ thuật Sliding Window, giúp cuộc trò chuyện diễn ra tự nhiên theo ngữ cảnh mà không bị tốn quá nhiều Token.', style='List Number')
doc.add_paragraph('2. RAG (Retrieval-Augmented Generation) & Đọc PDF: AI không chỉ chém gió chay mà còn có khả năng đọc hiểu file PDF (như Bảng điểm của sinh viên). File PDF được băm nhỏ (chunking), chuyển thành Vector (Embedding) và lưu vào Database PostgreSQL (pgvector). Khi người dùng hỏi, AI sẽ tìm kiếm các đoạn text liên quan trong DB để trả lời.', style='List Number')
doc.add_paragraph('3. Function Calling (Tools): Đây là tính năng "bá đạo" nhất. AI được trang bị công cụ JobMarketTool. Thay vì chỉ biết lý thuyết, khi người dùng hỏi về mức lương hay công việc, AI có khả năng tự động kích hoạt Tool để chạy xuống Database thật (nơi chứa dữ liệu cào từ TopCV/LinkedIn do bạn của bạn viết) để lấy dữ liệu thực tế và báo cáo cho người dùng.', style='List Number')

# Section 3
doc.add_heading('3. Cách làm & Các đoạn Code quan trọng', level=1)

doc.add_heading('A. Khởi tạo ChatClient và gắn Bộ Nhớ (Memory)', level=2)
doc.add_paragraph('Chúng ta thiết lập ChatClient để nó đóng vai trò là một cố vấn nghề nghiệp. Kỹ thuật MessageChatMemoryAdvisor được áp dụng để duy trì ngữ cảnh.')
doc.add_paragraph('File quan trọng: VirtualMentorService.java', style='Intense Quote')
p1 = doc.add_paragraph('// Cấu hình AI với System Prompt và Memory\nthis.chatClient = chatClientBuilder\n    .defaultSystem("""\n        ## IDENTITY & SCOPE\n        - You advise IT students on careers, skills, roadmaps, etc.\n        - CRITICAL LANGUAGE RULE: You MUST ALWAYS reply in the EXACT SAME LANGUAGE the user uses.\n    """)\n    .defaultAdvisors(\n        new MessageChatMemoryAdvisor(chatMemory, "default", 10)\n    )\n    .build();')
p1.style = 'Quote'

doc.add_heading('B. Xử lý RAG (Đọc PDF và Vector Database)', level=2)
doc.add_paragraph('Khi người dùng tải file PDF lên, hệ thống sẽ dùng PdfToMarkdownService để bóc tách chữ. Sau đó, nội dung này được gửi cho Spring AI để chuyển thành Vector.')
doc.add_paragraph('File quan trọng: VirtualMentorService.java', style='Quote')
p2 = doc.add_paragraph('// Đọc file PDF, thêm Metadata và lưu vào Vector Store\nDocument document = new Document(markdownContent, \n    Map.of("session_id", session.getSessionId(), "source", pdfUrl));\nvectorStore.add(List.of(document));\n\n// Khi Chat, gắn thêm QuestionAnswerAdvisor để AI tự đi tìm dữ liệu\nchatClient.prompt()\n    .advisors(new QuestionAnswerAdvisor(vectorStore, SearchRequest.defaults(), "... prompt ..."))\n    .stream().content();')
p2.style = 'Quote'

doc.add_heading('C. Function Calling (Trang bị Tool "Tay chân" cho AI)', level=2)
doc.add_paragraph('Đây là cách chúng ta cho AI gọi hàm Java. Ta định nghĩa một Function Bean, sau đó đăng ký nó vào ChatClient.')
doc.add_paragraph('File quan trọng: JobMarketTool.java', style='Quote')
p3 = doc.add_paragraph('@Service("jobMarketTool")\n@Description("Search for real-time IT jobs in Vietnam by keyword to get salary ranges.")\npublic class JobMarketTool implements Function<JobMarketTool.Request, JobMarketTool.Response> {\n    \n    private final RecruitmentRepository recruitmentRepository;\n\n    @Override\n    public Response apply(Request request) {\n        // Query database thật từ bảng recruitments\n        List<Recruitment> jobs = recruitmentRepository.findTop10ByTitleContainingIgnoreCase(request.keyword());\n        return new Response(jobs, "Found jobs...");\n    }\n}')
p3.style = 'Quote'
doc.add_paragraph('Và khai báo công cụ này cho AI biết mỗi khi Chat:')
p4 = doc.add_paragraph('chatClient.prompt()\n    .functions("jobMarketTool") // AI sẽ tự quyết định lúc nào cần gọi hàm này\n    .stream()')
p4.style = 'Quote'

# Section 4
doc.add_heading('4. Tổng Kết', level=1)
doc.add_paragraph('Bằng việc ứng dụng Spring AI, InteliPath đã chuyển mình từ một ứng dụng CRUD thông thường thành một hệ thống AI-Driven hiện đại. AI không chỉ đóng vai trò "Chatbot" mà đã thực sự trở thành một Agent (Trợ lý tự chủ) có khả năng đọc tài liệu riêng của người dùng (RAG) và tương tác trực tiếp với cơ sở dữ liệu của hệ thống thông qua Tools.')

doc.save('D:/Project/IntelIRoadMap/BaoCao_SpringAI.docx')
print('Tạo file BaoCao_SpringAI.docx thành công!')
